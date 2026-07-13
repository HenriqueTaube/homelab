from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Cores ──────────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1B, 0x3A, 0x5C)
STEEL  = RGBColor(0x2E, 0x6D, 0xA8)
LIGHT  = RGBColor(0xE8, 0xF1, 0xFA)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GRAY   = RGBColor(0x55, 0x55, 0x55)
DARK   = RGBColor(0x22, 0x22, 0x22)

# ── Helpers ────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def no_border(table):
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'none')
        el.set(qn('w:sz'),    '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    tblPr.append(tblBorders)

def fix_col_widths(table, widths_dxa):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'),    str(sum(widths_dxa)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.insert(0, tblW)
    for old in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(old)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    for old in tbl.findall(qn('w:tblGrid')):
        tbl.remove(old)
    tblGrid = OxmlElement('w:tblGrid')
    for w in widths_dxa:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(w))
        tblGrid.append(gridCol)
    tblPr.addnext(tblGrid)
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for old in tcPr.findall(qn('w:tcW')):
                tcPr.remove(old)
            tcW  = OxmlElement('w:tcW')
            tcW.set(qn('w:w'),    str(widths_dxa[ci]))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.insert(0, tcW)

def add_heading(doc, text, level=1):
    p   = doc.add_paragraph()
    pf  = p.paragraph_format
    pf.space_before = Pt(10 if level==1 else 6)
    pf.space_after  = Pt(4)
    run = p.add_run(text.upper() if level==1 else text)
    run.font.name   = 'Calibri'
    run.font.size   = Pt(13 if level==1 else 11)
    run.font.bold   = True
    run.font.color.rgb = NAVY
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    btm  = OxmlElement('w:bottom')
    btm.set(qn('w:val'),   'single')
    btm.set(qn('w:sz'),    '6')
    btm.set(qn('w:color'), '1B3A5C')
    pBdr.append(btm)
    pPr.append(pBdr)
    return p

def add_body(doc, text, bold=False, bullet=False):
    p   = doc.add_paragraph()
    pf  = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(3)
    if bullet:
        pf.left_indent = Cm(0.5)
        run = p.add_run('• ' + text)
    else:
        run = p.add_run(text)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(10.5)
    run.font.bold  = bold
    run.font.color.rgb = DARK
    return p

def add_kv_table(doc, rows, col_widths_dxa=(3575, 6400)):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    no_border(table)
    fix_col_widths(table, list(col_widths_dxa))
    for i, (key, val) in enumerate(rows):
        row = table.rows[i]
        bg = 'E8F1FA' if i % 2 == 0 else 'FFFFFF'
        for ci, (cell, txt, bld) in enumerate([(row.cells[0], key, True), (row.cells[1], val, False)]):
            set_cell_bg(cell, bg)
            p  = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent  = Cm(0.2)
            run = p.add_run(txt)
            run.font.name  = 'Calibri'
            run.font.size  = Pt(10)
            run.font.bold  = bld
            run.font.color.rgb = NAVY if bld else DARK
    return table

def add_scope_table(doc, system_name, components):
    tbl = doc.add_table(rows=1 + len(components), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    no_border(tbl)

    # Linha pai
    hrow = tbl.rows[0]
    hrow.cells[0].width = Cm(1.5)
    hrow.cells[1].width = Cm(16.0)
    for ci, txt in enumerate(['QTD', system_name]):
        set_cell_bg(hrow.cells[ci], '1B3A5C')
        p = hrow.cells[ci].paragraphs[0]
        p.paragraph_format.left_indent  = Cm(0.3)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(txt)
        r.font.name  = 'Calibri'
        r.font.size  = Pt(10)
        r.font.bold  = True
        r.font.color.rgb = WHITE
        if ci == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Linhas filhas
    for ri, (qtd, desc) in enumerate(components):
        row = tbl.rows[ri + 1]
        bg  = 'E8F1FA' if ri % 2 == 0 else 'FFFFFF'
        set_cell_bg(row.cells[0], bg)
        set_cell_bg(row.cells[1], bg)
        row.cells[0].width = Cm(1.5)
        row.cells[1].width = Cm(16.0)

        pq = row.cells[0].paragraphs[0]
        pq.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pq.paragraph_format.space_before = Pt(4)
        pq.paragraph_format.space_after  = Pt(4)
        rq = pq.add_run(qtd)
        rq.font.name  = 'Calibri'
        rq.font.size  = Pt(11)
        rq.font.bold  = True
        rq.font.color.rgb = NAVY

        pd = row.cells[1].paragraphs[0]
        pd.paragraph_format.left_indent  = Cm(0.3)
        pd.paragraph_format.space_before = Pt(4)
        pd.paragraph_format.space_after  = Pt(4)
        rd = pd.add_run(desc)
        rd.font.name  = 'Calibri'
        rd.font.size  = Pt(10)
        rd.font.color.rgb = DARK

    doc.add_paragraph()

def add_investment_block(doc, label, items, total_label, total_value):
    # items: (desc, qty, valor_unit, total)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(label)
    r.font.name  = 'Calibri'
    r.font.size  = Pt(11)
    r.font.bold  = True
    r.font.color.rgb = NAVY

    table = doc.add_table(rows=len(items) + 1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    no_border(table)
    fix_col_widths(table, [5200, 900, 2000, 1875])

    headers = ['Descrição', 'Qtd.', 'Valor Unit.', 'Total']
    aligns  = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
               WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]
    hrow = table.rows[0]
    for ci, h in enumerate(headers):
        set_cell_bg(hrow.cells[ci], '1B3A5C')
        p2 = hrow.cells[ci].paragraphs[0]
        p2.alignment = aligns[ci]
        p2.paragraph_format.left_indent  = Cm(0.2)
        p2.paragraph_format.space_before = Pt(3)
        p2.paragraph_format.space_after  = Pt(3)
        run = p2.add_run(h)
        run.font.name  = 'Calibri'
        run.font.size  = Pt(10)
        run.font.bold  = True
        run.font.color.rgb = WHITE

    for ri, (desc, qty, valor_unit, total) in enumerate(items):
        row = table.rows[ri + 1]
        bg  = 'E8F1FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, (txt, align) in enumerate(zip([desc, qty, valor_unit, total], aligns)):
            set_cell_bg(row.cells[ci], bg)
            p3 = row.cells[ci].paragraphs[0]
            p3.alignment = align
            p3.paragraph_format.left_indent  = Cm(0.2)
            p3.paragraph_format.space_before = Pt(3)
            p3.paragraph_format.space_after  = Pt(3)
            run = p3.add_run(txt)
            run.font.name  = 'Calibri'
            run.font.size  = Pt(10)
            run.font.color.rgb = DARK

    total_row = table.add_row()
    for ci in range(4):
        set_cell_bg(total_row.cells[ci], '1B3A5C')
        p4 = total_row.cells[ci].paragraphs[0]
        p4.paragraph_format.left_indent  = Cm(0.2)
        p4.paragraph_format.space_before = Pt(3)
        p4.paragraph_format.space_after  = Pt(3)
        txt = total_label if ci == 0 else (total_value if ci == 3 else '')
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 3 else WD_ALIGN_PARAGRAPH.LEFT
        run = p4.add_run(txt)
        run.font.name  = 'Calibri'
        run.font.size  = Pt(10)
        run.font.bold  = True
        run.font.color.rgb = WHITE

def build_footer(doc):
    sec = doc.sections[0]
    sec.footer_distance = Cm(0.8)
    footer = sec.footer
    footer.is_linked_to_previous = False

    fp = footer.paragraphs[0]
    fp.clear()
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after  = Pt(0)

    pPr = fp._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    for val, pos in [('center', '4819'), ('right', '9638')]:
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), val)
        tab.set(qn('w:pos'), pos)
        tabs.append(tab)
    pPr.append(tabs)

    HYPERLINK_TYPE = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink'

    def make_rPr():
        rPr = OxmlElement('w:rPr')
        fn  = OxmlElement('w:rFonts')
        fn.set(qn('w:ascii'), 'Calibri'); fn.set(qn('w:hAnsi'), 'Calibri')
        sz  = OxmlElement('w:sz');  sz.set(qn('w:val'), '15')
        clr = OxmlElement('w:color'); clr.set(qn('w:val'), '1B3A5C')
        rPr.extend([fn, sz, clr])
        return rPr

    def add_text(text):
        r = OxmlElement('w:r')
        r.append(make_rPr())
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        fp._p.append(r)

    def add_field(code):
        for ftype in ('begin', 'instr', 'separate', 'end'):
            r = OxmlElement('w:r')
            r.append(make_rPr())
            if ftype == 'instr':
                instr = OxmlElement('w:instrText')
                instr.set(qn('xml:space'), 'preserve')
                instr.text = f' {code} '
                r.append(instr)
            else:
                fc = OxmlElement('w:fldChar')
                fc.set(qn('w:fldCharType'), ftype if ftype != 'instr' else 'begin')
                r.append(fc)
            fp._p.append(r)

    def add_hyperlink(url, text):
        rel_id = footer.part.relate_to(url, HYPERLINK_TYPE, is_external=True)
        hl = OxmlElement('w:hyperlink')
        hl.set(qn('r:id'), rel_id)
        r = OxmlElement('w:r')
        r.append(make_rPr())
        t = OxmlElement('w:t'); t.text = text
        r.append(t)
        hl.append(r)
        fp._p.append(hl)

    # Esquerda: WhatsApp
    add_hyperlink('https://wa.me/5547991661269', '47 9 9166-1269')

    # Centro: PAGE / NUMPAGES
    add_text('\t')
    add_field('PAGE')
    add_text(' / ')
    add_field('NUMPAGES')

    # Direita: site
    add_text('\t')
    add_hyperlink('https://www.taubeequipamentos.com.br', 'taubeequipamentos.com.br')

# ── Documento ──────────────────────────────────────────────────────────────────
doc = Document()

sec = doc.sections[0]
sec.top_margin    = Cm(1.5)
sec.bottom_margin = Cm(2.0)
sec.left_margin   = Cm(2.0)
sec.right_margin  = Cm(2.0)

# ── CABEÇALHO ──────────────────────────────────────────────────────────────────
header_tbl = doc.add_table(rows=1, cols=2)
header_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
no_border(header_tbl)

logo_cell  = header_tbl.rows[0].cells[0]
title_cell = header_tbl.rows[0].cells[1]
logo_cell.width  = Cm(5.5)
title_cell.width = Cm(12.0)

set_cell_bg(logo_cell,  '1B3A5C')
set_cell_bg(title_cell, '1B3A5C')
logo_cell.vertical_alignment  = WD_ALIGN_VERTICAL.CENTER
title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

logo_p = logo_cell.paragraphs[0]
logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
logo_p.paragraph_format.space_before = Pt(6)
logo_p.paragraph_format.space_after  = Pt(6)
logo_p.paragraph_format.left_indent  = Cm(0.2)
run_logo = logo_p.add_run()
run_logo.add_picture('/home/henrique/projetos/editor/proposta/assets/logo taube minizado.jpg', height=Cm(2.8))

title_p = title_cell.paragraphs[0]
title_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
title_p.paragraph_format.space_before = Pt(10)
title_p.paragraph_format.space_after  = Pt(2)
title_p.paragraph_format.right_indent = Cm(0.3)
r1 = title_p.add_run('TAUBE EQUIPAMENTOS')
r1.font.name  = 'Calibri'
r1.font.size  = Pt(16)
r1.font.bold  = True
r1.font.color.rgb = WHITE

p2 = title_cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p2.paragraph_format.right_indent = Cm(0.3)
p2.paragraph_format.space_before = Pt(2)
p2.paragraph_format.space_after  = Pt(10)
r2 = p2.add_run('www.taubeequipamentos.com.br')
r2.font.name  = 'Calibri'
r2.font.size  = Pt(10)
r2.font.bold  = False
r2.font.color.rgb = RGBColor(0xA8, 0xC8, 0xE8)

# Linha separadora
sep = doc.add_paragraph()
sep.paragraph_format.space_before = Pt(0)
sep.paragraph_format.space_after  = Pt(0)
pPr  = sep._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
btm  = OxmlElement('w:bottom')
btm.set(qn('w:val'),   'single')
btm.set(qn('w:sz'),    '12')
btm.set(qn('w:color'), '2E6DA8')
pBdr.append(btm)
pPr.append(pBdr)

# Número da proposta
num_p = doc.add_paragraph()
num_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
num_p.paragraph_format.space_before = Pt(8)
num_p.paragraph_format.space_after  = Pt(8)
r_num = num_p.add_run('PROPOSTA COMERCIAL N° 078 / 2026 – SERRA INDÚSTRIA E COMÉRCIO DE PAPEIS LTDA.')
r_num.font.name  = 'Calibri'
r_num.font.size  = Pt(13)
r_num.font.bold  = True
r_num.font.color.rgb = NAVY

# ── DADOS DO DESTINATÁRIO ──────────────────────────────────────────────────────
dest_tbl = doc.add_table(rows=1, cols=2)
dest_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
no_border(dest_tbl)

left_cell  = dest_tbl.rows[0].cells[0]
right_cell = dest_tbl.rows[0].cells[1]
left_cell.width  = Cm(9.0)
right_cell.width = Cm(8.5)
set_cell_bg(left_cell,  'E8F1FA')
set_cell_bg(right_cell, 'FFFFFF')

lp = left_cell.paragraphs[0]
lp.paragraph_format.left_indent  = Cm(0.3)
lp.paragraph_format.space_before = Pt(4)
lp.paragraph_format.space_after  = Pt(2)
r = lp.add_run('DESTINATÁRIO')
r.font.name  = 'Calibri'
r.font.size  = Pt(8)
r.font.bold  = True
r.font.color.rgb = STEEL

for line in [
    ('Serra Indústria e Comércio de Papeis Ltda.', True),
    ('Aos cuidados de: Sr. Samuel', False),
    ('Departamento de Compras', False),
    ('Fone: 27 99727-3731', False),
    ('compras@serrapapeis.com.br', False),
]:
    p = left_cell.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(line[0])
    run.font.name  = 'Calibri'
    run.font.size  = Pt(11 if line[1] else 10)
    run.font.bold  = line[1]
    run.font.color.rgb = NAVY if line[1] else DARK

rp = right_cell.paragraphs[0]
rp.paragraph_format.left_indent  = Cm(0.3)
rp.paragraph_format.space_before = Pt(4)
rp.paragraph_format.space_after  = Pt(2)
r = rp.add_run('EMISSÃO')
r.font.name  = 'Calibri'
r.font.size  = Pt(8)
r.font.bold  = True
r.font.color.rgb = STEEL

for line in [
    'Data: 02/06/2026',
    'Validade: 15 dias a partir da emissão',
    'CNPJ: 11.478.834/0001-23',
    'IE: 256.047.162',
    'Simples Nacional',
]:
    p = right_cell.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(line)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(10)
    run.font.color.rgb = DARK

for cell in [left_cell, right_cell]:
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(0)

doc.add_paragraph()

# ── 1. APRESENTAÇÃO ────────────────────────────────────────────────────────────
add_heading(doc, '1. Apresentação')
add_body(doc,
    'A TAUBE EQUIPAMENTOS apresenta sua proposta para fornecimento de Apalpador Pneumático '
    'TB Axial, desenvolvido para aplicação em válvula apalpadora de vestimenta para mancais '
    'de dupla ação.')
doc.add_paragraph()

# ── 2. OBJETIVO ────────────────────────────────────────────────────────────────
add_heading(doc, '2. Objetivo')
add_body(doc,
    'A presente proposta contempla o fornecimento de 02 apalpadores pneumáticos TB Axial, '
    'projetados para proporcionar confiabilidade operacional, resistência mecânica e elevada '
    'durabilidade em ambiente industrial do setor papeleiro.')
# ── 3. ITEM 01 ─────────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '3. Item 01 – Apalpador Pneumático TB Axial')

add_heading(doc, '3.1 Quantidade', level=2)

itens_078 = [('02', 'Apalpador Pneumático TB Axial')]
escopo_tbl = doc.add_table(rows=1 + len(itens_078), cols=2)
escopo_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
no_border(escopo_tbl)
fix_col_widths(escopo_tbl, [1363, 8562])

hrow = escopo_tbl.rows[0]
for ci, txt in enumerate(['QTD', 'PRODUTO']):
    set_cell_bg(hrow.cells[ci], '1B3A5C')
    p = hrow.cells[ci].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(txt)
    r.font.name  = 'Calibri'
    r.font.size  = Pt(10)
    r.font.bold  = True
    r.font.color.rgb = WHITE

for ri, (qtd, produto) in enumerate(itens_078):
    row = escopo_tbl.rows[ri + 1]
    bg  = 'E8F1FA' if ri % 2 == 0 else 'FFFFFF'
    set_cell_bg(row.cells[0], bg)
    set_cell_bg(row.cells[1], bg)
    pq = row.cells[0].paragraphs[0]
    pq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pq.paragraph_format.space_before = Pt(4)
    pq.paragraph_format.space_after  = Pt(4)
    rq = pq.add_run(qtd)
    rq.font.name  = 'Calibri'
    rq.font.size  = Pt(11)
    rq.font.bold  = True
    rq.font.color.rgb = NAVY
    pd = row.cells[1].paragraphs[0]
    pd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pd.paragraph_format.space_before = Pt(4)
    pd.paragraph_format.space_after  = Pt(4)
    rd = pd.add_run(produto)
    rd.font.name  = 'Calibri'
    rd.font.size  = Pt(10)
    rd.font.color.rgb = DARK

doc.add_paragraph()

add_heading(doc, '3.2 Dados de Aplicação', level=2)
add_kv_table(doc, [
    ('Tipo de Operação',     'Mancais de Dupla Ação'),
    ('Material Estrutural',  'Aço Inoxidável AISI 304'),
    ('Material da Mola',     'Aço Inoxidável AISI 302'),
    ('Elementos de Vedação', 'Viton de alta resistência térmica'),
    ('Sistema de Retorno',   'Mola helicoidal ou opcional por contrapeso'),
    ('Ajustes Disponíveis',  'Regulagem Vertical e Horizontal'),
    ('Sistema de Acionamento', 'Pneumático'),
    ('Ambiente de Aplicação',  'Máquina de papel e celulose'),
    ('Resistência à Corrosão', 'Elevada, adequada para ambientes úmidos e agressivos'),
])

doc.add_paragraph()
add_heading(doc, '3.3 Características Construtivas', level=2)
add_kv_table(doc, [
    ('Estrutura Principal',
     'O conjunto é fabricado integralmente em aço inoxidável AISI 304, proporcionando elevada '
     'resistência mecânica, excelente resistência à corrosão e longa vida útil mesmo em '
     'ambientes industriais severos do setor papeleiro.'),
    ('Hastes de Posicionamento',
     'Construídas em perfil tubular de aço inoxidável AISI 304, permitindo ajustes verticais '
     'e horizontais para otimização do posicionamento operacional e adequação às condições '
     'específicas de instalação.'),
    ('Base de Fixação',
     'Fabricada em aço inoxidável AISI 304, desenvolvida para proporcionar elevada rigidez '
     'estrutural, estabilidade operacional e facilidade de montagem.'),
    ('Corpo do Apalpador',
     'Construído totalmente em aço inoxidável, com parafusos, elementos de fixação e sistema '
     'de ajuste de tensão da mola também em inox, garantindo durabilidade superior e reduzindo '
     'necessidades de manutenção corretiva.'),
    ('Palm Guide',
     'Elemento de contato com a vestimenta fabricado com superfície cerâmica de alta resistência '
     'ao desgaste, proporcionando baixo atrito, maior estabilidade operacional e redução do '
     'desgaste da vestimenta. O retorno do conjunto pode ser realizado por mola helicoidal ou, '
     'opcionalmente, por sistema de contrapeso, conforme necessidade operacional do cliente.'),
    ('Sistema de Vedação',
     'Projeto com vedação interna integral, minimizando perdas pneumáticas e eliminando '
     'vazamentos durante a operação. Todos os elementos de vedação são fabricados em Viton, '
     'material com elevada resistência química e térmica, garantindo maior confiabilidade e '
     'vida útil do equipamento.'),
])

doc.add_paragraph()
add_heading(doc, '3.4 Recursos Operacionais', level=2)
for item in [
    'Ajuste vertical e horizontal do conjunto',
    'Regulagem da tensão da mola',
    'Sistema pneumático de elevada confiabilidade',
    'Vedação interna de alta eficiência',
    'Construção integral em aço inoxidável',
    'Baixa necessidade de manutenção',
    'Fácil instalação e substituição de componentes',
    'Elevada resistência à corrosão e ao desgaste',
]:
    add_body(doc, item, bullet=True)

doc.add_paragraph()
add_heading(doc, '3.5 Benefícios Operacionais', level=2)
for item in [
    'Maior estabilidade no controle da vestimenta',
    'Redução de intervenções de manutenção',
    'Menor desgaste dos componentes de contato',
    'Maior vida útil da vestimenta',
    'Elevada confiabilidade operacional',
    'Excelente desempenho em ambientes úmidos e agressivos',
    'Facilidade de ajuste durante paradas programadas',
    'Redução do custo total de manutenção',
]:
    add_body(doc, item, bullet=True)

# ── 4. INVESTIMENTO ────────────────────────────────────────────────────────────
doc.add_page_break()

# Imagem do produto
img_p = doc.add_paragraph()
img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
img_p.paragraph_format.space_before = Pt(6)
img_p.paragraph_format.space_after  = Pt(10)
img_p.add_run().add_picture('/tmp/apalpador_img.jpg', height=Cm(9))

add_heading(doc, '4. Investimento')
add_investment_block(
    doc,
    'ITEM 01 – APALPADOR PNEUMÁTICO TB AXIAL',
    [
        ('Apalpador Pneumático TB Axial – Palm Guide Inox', '02', 'R$ 11.000,00', 'R$ 22.000,00'),
    ],
    'TOTAL',
    'R$ 22.000,00'
)

doc.add_paragraph()

# ── 5. PRAZO DE ENTREGA ────────────────────────────────────────────────────────
add_heading(doc, '5. Prazo de Entrega')
add_body(doc, 'Item 01:  15 dias corridos após confirmação do pedido.')
doc.add_paragraph()

# ── 6. CONDIÇÕES DE PAGAMENTO ──────────────────────────────────────────────────
add_heading(doc, '6. Condições de Pagamento')
add_body(doc, 'Item 01:', bold=True)
add_body(doc, '28/45 DDL', bullet=True)
doc.add_paragraph()

# ── 7. EXCLUSÕES DO FORNECIMENTO ──────────────────────────────────────────────
add_heading(doc, '7. Exclusões do Fornecimento')
add_body(doc, 'Não estão inclusos:')
for item in [
    'Transporte', 'Seguro', 'Armazenagem', 'Montagem mecânica',
    'Instalação elétrica', 'Startup', 'Supervisão de montagem',
    'Infraestrutura hidráulica', 'Infraestrutura elétrica',
    'Equipamentos de içamento', 'Peças sobressalentes',
    'Integrações não descritas nesta proposta',
]:
    add_body(doc, item, bullet=True)
doc.add_paragraph()

# ── 8. GARANTIA ────────────────────────────────────────────────────────────────
add_heading(doc, '8. Garantia')
add_body(doc,
    'Garantia de 180 dias após startup ou 12 meses após faturamento, '
    'prevalecendo o primeiro evento ocorrido.')
add_body(doc,
    'A garantia cobre defeitos comprovados de fabricação, materiais e montagem, desde que os '
    'equipamentos sejam transportados, armazenados, instalados e operados conforme as '
    'recomendações técnicas fornecidas.')
doc.add_paragraph()

# ── 9. VALIDADE ────────────────────────────────────────────────────────────────
add_heading(doc, '9. Validade da Proposta')
add_body(doc, 'Esta proposta possui validade de 15 dias a partir da data de emissão.')

# ── RODAPÉ ─────────────────────────────────────────────────────────────────────
doc.add_paragraph()
sep2 = doc.add_paragraph()
sep2.paragraph_format.space_before = Pt(0)
sep2.paragraph_format.space_after  = Pt(6)
pPr2  = sep2._p.get_or_add_pPr()
pBdr2 = OxmlElement('w:pBdr')
top2  = OxmlElement('w:top')
top2.set(qn('w:val'),   'single')
top2.set(qn('w:sz'),    '6')
top2.set(qn('w:color'), '1B3A5C')
pBdr2.append(top2)
pPr2.append(pBdr2)

footer_tbl = doc.add_table(rows=1, cols=2)
footer_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
no_border(footer_tbl)
fl = footer_tbl.rows[0].cells[0]
fr = footer_tbl.rows[0].cells[1]
fl.width = Cm(11)
fr.width = Cm(6.5)

HYPERLINK_TYPE = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink'

def _gray(para, text, size=9):
    r = para.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(size); r.font.color.rgb = GRAY

def _body_link(para, url, text, size=9):
    rel_id = doc.part.relate_to(url, HYPERLINK_TYPE, is_external=True)
    hl = OxmlElement('w:hyperlink'); hl.set(qn('r:id'), rel_id)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    fn = OxmlElement('w:rFonts')
    fn.set(qn('w:ascii'), 'Calibri'); fn.set(qn('w:hAnsi'), 'Calibri')
    sz  = OxmlElement('w:sz');    sz.set(qn('w:val'), str(size * 2))
    clr = OxmlElement('w:color'); clr.set(qn('w:val'), '555555')
    rPr.extend([fn, sz, clr]); r.append(rPr)
    t = OxmlElement('w:t'); t.text = text; r.append(t)
    hl.append(r); para._p.append(hl)

def _rp(cell, first=False):
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    return p

# Célula esquerda
p_l = fl.paragraphs[0]
p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_l.paragraph_format.space_before = Pt(0)
p_l.paragraph_format.space_after  = Pt(0)
_gray(p_l, 'Aproveitamos para nos colocar à disposição para quaisquer esclarecimentos adicionais.\nAtenciosamente,\n\nTAUBE EQUIPAMENTOS')

# Célula direita — um parágrafo por linha para permitir hyperlinks
_gray(_rp(fr, first=True), 'CNPJ: 11.478.834/0001-23 | IE: 256.047.162')
_gray(_rp(fr), 'Simples Nacional – Não gera crédito de ICMS/IPI')
_body_link(_rp(fr), 'https://www.taubeequipamentos.com.br', 'www.taubeequipamentos.com.br')
p_tel = _rp(fr)
_gray(p_tel, '47 3234-1222 | ')
_body_link(p_tel, 'https://wa.me/5547991661269', '47 9 9166-1269')

build_footer(doc)

out = '/home/henrique/projetos/editor/proposta/PROPOSTA COMERCIAL 078 2026 SERRA PAPEIS - TAUBE.docx'
doc.save(out)
print(f'Salvo em: {out}')
