from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Cores ──────────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1B, 0x3A, 0x5C)   # cabeçalho / títulos
STEEL  = RGBColor(0x2E, 0x6D, 0xA8)   # subtítulos / linhas de tabela
LIGHT  = RGBColor(0xE8, 0xF1, 0xFA)   # fundo linha de cabeçalho de tabela
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GRAY   = RGBColor(0x55, 0x55, 0x55)
DARK   = RGBColor(0x22, 0x22, 0x22)

# ── Helpers ────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),   val.get('val','single'))
            el.set(qn('w:sz'),    val.get('sz','4'))
            el.set(qn('w:color'), val.get('color','auto'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def no_border(table):
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'none')
        el.set(qn('w:sz'),    '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    tblPr.append(tblBorders)

def fix_col_widths(table, widths_dxa):
    """Força larguras fixas em twips substituindo os elementos existentes."""
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    # Substituir tblW existente
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'),    str(sum(widths_dxa)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.insert(0, tblW)
    # Substituir tblLayout existente
    for old in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(old)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    # Substituir tblGrid existente
    for old in tbl.findall(qn('w:tblGrid')):
        tbl.remove(old)
    tblGrid = OxmlElement('w:tblGrid')
    for w in widths_dxa:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(w))
        tblGrid.append(gridCol)
    tblPr.addnext(tblGrid)
    # Largura em cada célula
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for old in tcPr.findall(qn('w:tcW')):
                tcPr.remove(old)
            tcW  = OxmlElement('w:tcW')
            tcW.set(qn('w:w'),    str(widths_dxa[ci]))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.insert(0, tcW)

def para_fmt(para, font_name='Calibri', size=11, bold=False,
             color=None, align=None, space_before=0, space_after=4):
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    for run in para.runs:
        run.font.name  = font_name
        run.font.size  = Pt(size)
        run.font.bold  = bold
        if color:
            run.font.color.rgb = color

def add_heading(doc, text, level=1):
    p   = doc.add_paragraph()
    pf  = p.paragraph_format
    pf.space_before = Pt(10 if level==1 else 6)
    pf.space_after  = Pt(4)
    run = p.add_run(text.upper() if level==1 else text)
    run.font.name   = 'Calibri'
    run.font.size   = Pt(13 if level==1 else 11)
    run.font.bold   = True
    run.font.color.rgb = NAVY
    # linha colorida abaixo
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    btm  = OxmlElement('w:bottom')
    btm.set(qn('w:val'),   'single')
    btm.set(qn('w:sz'),    '6')
    btm.set(qn('w:color'), '1B3A5C')
    pBdr.append(btm)
    pPr.append(pBdr)
    return p

def add_body(doc, text, bold=False, bullet=False, indent=False):
    p   = doc.add_paragraph()
    pf  = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(3)
    if bullet:
        pf.left_indent = Cm(0.5)
        run = p.add_run('• ' + text)
    elif indent:
        pf.left_indent = Cm(0.5)
        run = p.add_run(text)
    else:
        run = p.add_run(text)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(10.5)
    run.font.bold  = bold
    run.font.color.rgb = DARK
    return p

def add_kv_table(doc, rows, col_widths_dxa=(3575, 6400)):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    no_border(table)
    fix_col_widths(table, list(col_widths_dxa))
    for i, (key, val) in enumerate(rows):
        row = table.rows[i]
        bg = 'E8F1FA' if i % 2 == 0 else 'FFFFFF'
        for ci, (cell, txt, bld) in enumerate([(row.cells[0], key, True), (row.cells[1], val, False)]):
            set_cell_bg(cell, bg)
            p  = cell.paragraphs[0]
            pf = p.paragraph_format
            pf.space_before = Pt(3)
            pf.space_after  = Pt(3)
            pf.left_indent  = Cm(0.2)
            run = p.add_run(txt)
            run.font.name  = 'Calibri'
            run.font.size  = Pt(10)
            run.font.bold  = bld
            run.font.color.rgb = NAVY if bld else DARK
    return table

def add_investment_block(doc, label, items, total_label, total_value):
    # items: (desc, qty, valor_unit, total)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(label)
    r.font.name  = 'Calibri'
    r.font.size  = Pt(11)
    r.font.bold  = True
    r.font.color.rgb = NAVY

    table = doc.add_table(rows=len(items)+1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    no_border(table)
    fix_col_widths(table, [5200, 900, 2000, 1875])

    headers = ['Descrição', 'Qtd.', 'Valor Unit.', 'Total']
    aligns  = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
               WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]
    hrow = table.rows[0]
    for ci, h in enumerate(headers):
        set_cell_bg(hrow.cells[ci], '1B3A5C')
        p2 = hrow.cells[ci].paragraphs[0]
        p2.alignment = aligns[ci]
        p2.paragraph_format.left_indent  = Cm(0.2)
        p2.paragraph_format.space_before = Pt(3)
        p2.paragraph_format.space_after  = Pt(3)
        run = p2.add_run(h)
        run.font.name  = 'Calibri'
        run.font.size  = Pt(10)
        run.font.bold  = True
        run.font.color.rgb = WHITE

    for ri, (desc, qty, valor_unit, total) in enumerate(items):
        row = table.rows[ri+1]
        bg  = 'E8F1FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, (txt, align) in enumerate(zip([desc, qty, valor_unit, total], aligns)):
            set_cell_bg(row.cells[ci], bg)
            p3 = row.cells[ci].paragraphs[0]
            p3.alignment = align
            p3.paragraph_format.left_indent  = Cm(0.2)
            p3.paragraph_format.space_before = Pt(3)
            p3.paragraph_format.space_after  = Pt(3)
            run = p3.add_run(txt)
            run.font.name  = 'Calibri'
            run.font.size  = Pt(10)
            run.font.color.rgb = DARK

    total_row = table.add_row()
    for ci in range(4):
        set_cell_bg(total_row.cells[ci], '1B3A5C')
        p4 = total_row.cells[ci].paragraphs[0]
        p4.paragraph_format.left_indent  = Cm(0.2)
        p4.paragraph_format.space_before = Pt(3)
        p4.paragraph_format.space_after  = Pt(3)
        txt = total_label if ci == 0 else (total_value if ci == 3 else '')
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 3 else WD_ALIGN_PARAGRAPH.LEFT
        run = p4.add_run(txt)
        run.font.name  = 'Calibri'
        run.font.size  = Pt(10)
        run.font.bold  = True
        run.font.color.rgb = WHITE

    return table

def build_footer(doc):
    sec = doc.sections[0]
    sec.footer_distance = Cm(0.8)
    footer = sec.footer
    footer.is_linked_to_previous = False

    fp = footer.paragraphs[0]
    fp.clear()
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after  = Pt(0)

    pPr = fp._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    for val, pos in [('center', '4819'), ('right', '9638')]:
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), val)
        tab.set(qn('w:pos'), pos)
        tabs.append(tab)
    pPr.append(tabs)

    HYPERLINK_TYPE = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink'

    def make_rPr():
        rPr = OxmlElement('w:rPr')
        fn  = OxmlElement('w:rFonts')
        fn.set(qn('w:ascii'), 'Calibri'); fn.set(qn('w:hAnsi'), 'Calibri')
        sz  = OxmlElement('w:sz');  sz.set(qn('w:val'), '15')
        clr = OxmlElement('w:color'); clr.set(qn('w:val'), '1B3A5C')
        rPr.extend([fn, sz, clr])
        return rPr

    def add_text(text):
        r = OxmlElement('w:r')
        r.append(make_rPr())
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        fp._p.append(r)

    def add_field(code):
        for ftype in ('begin', 'instr', 'separate', 'end'):
            r = OxmlElement('w:r')
            r.append(make_rPr())
            if ftype == 'instr':
                instr = OxmlElement('w:instrText')
                instr.set(qn('xml:space'), 'preserve')
                instr.text = f' {code} '
                r.append(instr)
            else:
                fc = OxmlElement('w:fldChar')
                fc.set(qn('w:fldCharType'), ftype if ftype != 'instr' else 'begin')
                r.append(fc)
            fp._p.append(r)

    def add_hyperlink(url, text):
        rel_id = footer.part.relate_to(url, HYPERLINK_TYPE, is_external=True)
        hl = OxmlElement('w:hyperlink')
        hl.set(qn('r:id'), rel_id)
        r = OxmlElement('w:r')
        r.append(make_rPr())
        t = OxmlElement('w:t'); t.text = text
        r.append(t)
        hl.append(r)
        fp._p.append(hl)

    # Esquerda: WhatsApp
    add_hyperlink('https://wa.me/5547991661269', '47 9 9166-1269')

    # Centro: PAGE / NUMPAGES
    add_text('\t')
    add_field('PAGE')
    add_text(' / ')
    add_field('NUMPAGES')

    # Direita: site
    add_text('\t')
    add_hyperlink('https://www.taubeequipamentos.com.br', 'taubeequipamentos.com.br')

# ── Documento ──────────────────────────────────────────────────────────────────
doc = Document()

# Margens
sec = doc.sections[0]
sec.top_margin    = Cm(1.5)
sec.bottom_margin = Cm(2.0)
sec.left_margin   = Cm(2.0)
sec.right_margin  = Cm(2.0)

# ── CABEÇALHO ──────────────────────────────────────────────────────────────────
header_tbl = doc.add_table(rows=1, cols=2)
header_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
no_border(header_tbl)

logo_cell  = header_tbl.rows[0].cells[0]
title_cell = header_tbl.rows[0].cells[1]

logo_cell.width  = Cm(5.5)
title_cell.width = Cm(12.0)

set_cell_bg(logo_cell,  '1B3A5C')
set_cell_bg(title_cell, '1B3A5C')

logo_cell.vertical_alignment  = WD_ALIGN_VERTICAL.CENTER
title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# Logo
logo_p = logo_cell.paragraphs[0]
logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
logo_p.paragraph_format.space_before = Pt(6)
logo_p.paragraph_format.space_after  = Pt(6)
logo_p.paragraph_format.left_indent  = Cm(0.2)
run_logo = logo_p.add_run()
run_logo.add_picture('/home/henrique/projetos/editor/proposta/assets/logo taube minizado.jpg', height=Cm(2.8))

# Nome da empresa no cabeçalho
title_p = title_cell.paragraphs[0]
title_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
title_p.paragraph_format.space_before = Pt(10)
title_p.paragraph_format.space_after  = Pt(2)
title_p.paragraph_format.right_indent = Cm(0.3)
r1 = title_p.add_run('TAUBE EQUIPAMENTOS')
r1.font.name  = 'Calibri'
r1.font.size  = Pt(16)
r1.font.bold  = True
r1.font.color.rgb = WHITE

p2 = title_cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p2.paragraph_format.right_indent = Cm(0.3)
p2.paragraph_format.space_before = Pt(2)
p2.paragraph_format.space_after  = Pt(10)
r2 = p2.add_run('www.taubeequipamentos.com.br')
r2.font.name  = 'Calibri'
r2.font.size  = Pt(10)
r2.font.bold  = False
r2.font.color.rgb = RGBColor(0xA8, 0xC8, 0xE8)

# Linha separadora
sep = doc.add_paragraph()
sep.paragraph_format.space_before = Pt(0)
sep.paragraph_format.space_after  = Pt(0)
pPr  = sep._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
btm  = OxmlElement('w:bottom')
btm.set(qn('w:val'),   'single')
btm.set(qn('w:sz'),    '12')
btm.set(qn('w:color'), '2E6DA8')
pBdr.append(btm)
pPr.append(pBdr)

# Número da proposta abaixo do cabeçalho
num_p = doc.add_paragraph()
num_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
num_p.paragraph_format.space_before = Pt(8)
num_p.paragraph_format.space_after  = Pt(8)
r_num = num_p.add_run('PROPOSTA COMERCIAL N° 081 / 2026 – IRANI PAPEL E EMBALAGENS S.A.')
r_num.font.name  = 'Calibri'
r_num.font.size  = Pt(13)
r_num.font.bold  = True
r_num.font.color.rgb = NAVY

# ── DADOS DO DESTINATÁRIO ──────────────────────────────────────────────────────


dest_tbl = doc.add_table(rows=1, cols=2)
dest_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
no_border(dest_tbl)

left_cell  = dest_tbl.rows[0].cells[0]
right_cell = dest_tbl.rows[0].cells[1]
left_cell.width  = Cm(9.0)
right_cell.width = Cm(8.5)

set_cell_bg(left_cell,  'E8F1FA')
set_cell_bg(right_cell, 'FFFFFF')

# Destinatário
lp = left_cell.paragraphs[0]
lp.paragraph_format.left_indent  = Cm(0.3)
lp.paragraph_format.space_before = Pt(4)
lp.paragraph_format.space_after  = Pt(2)
r = lp.add_run('DESTINATÁRIO')
r.font.name  = 'Calibri'
r.font.size  = Pt(8)
r.font.bold  = True
r.font.color.rgb = STEEL

for line in [
    ('Irani Papel e Embalagens S.A.', True),
    ('Vargem Bonita – SC  |  CEP: 89600-000', False),
    ('Fone: (49) 3548-9100', False),
    ('Att. Diogo Kleimpaul', False),
    ('diogo.kleimpaul@irani.com.br', False),
]:
    p = left_cell.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(line[0])
    run.font.name  = 'Calibri'
    run.font.size  = Pt(10 if not line[1] else 11)
    run.font.bold  = line[1]
    run.font.color.rgb = NAVY if line[1] else DARK

# Emitente
rp = right_cell.paragraphs[0]
rp.paragraph_format.left_indent  = Cm(0.3)
rp.paragraph_format.space_before = Pt(4)
rp.paragraph_format.space_after  = Pt(2)
r = rp.add_run('EMISSÃO')
r.font.name  = 'Calibri'
r.font.size  = Pt(8)
r.font.bold  = True
r.font.color.rgb = STEEL

for line in [
    ('Data: 08/06/2026', False),
    ('Validade: 15 dias a partir da emissão', False),
    ('CNPJ: 11.478.834/0001-23', False),
    ('IE: 256.047.162', False),
    ('Simples Nacional', False),
]:
    p = right_cell.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(line[0])
    run.font.name  = 'Calibri'
    run.font.size  = Pt(10)
    run.font.bold  = line[1]
    run.font.color.rgb = DARK

# Padding final nas células
for cell in [left_cell, right_cell]:
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(0)

doc.add_paragraph()

# ── 1. APRESENTAÇÃO ────────────────────────────────────────────────────────────
add_heading(doc, '1. Apresentação')
add_body(doc,
    'A TAUBE EQUIPAMENTOS apresenta sua proposta para fornecimento de Sistema de Corte de '
    'Refilo por Jato de Água e Skid de Pressurização, desenvolvidos com tecnologia própria '
    'e fabricação nacional.')
add_body(doc,
    'A solução proposta visa aumentar a confiabilidade operacional do processo, reduzir perdas '
    'por quebra de folha e proporcionar maior facilidade de manutenção, contribuindo diretamente '
    'para a eficiência da linha de produção.')
doc.add_paragraph()

# ── 2. OBJETIVO ────────────────────────────────────────────────────────────────
add_heading(doc, '2. Objetivo')
add_body(doc,
    'A presente proposta contempla o fornecimento de Sistema de Corte de Refilo por Jato de Água '
    'de Alta Pressão e Skid de Bombeamento, para aplicação na linha de produção da Irani Papel '
    'e Embalagens S.A. – Unidade Vargem Bonita/SC.')
add_body(doc,
    'O sistema foi desenvolvido para proporcionar elevada estabilidade operacional, precisão no '
    'corte de refilo, redução de perdas e maior facilidade de operação e manutenção.')
# ── 3. ITEM 01 ─────────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '3. Item 01 – Sistema de Corte de Refilo com Microposicionamento')

add_heading(doc, '3.1 Quantidade', level=2)

itens_081 = [
    ('02', 'Corpos de Corte'),
    ('02', 'Suportes de Fixação'),
]
escopo_tbl = doc.add_table(rows=1 + len(itens_081), cols=2)
escopo_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
no_border(escopo_tbl)
fix_col_widths(escopo_tbl, [1363, 8562])

hrow = escopo_tbl.rows[0]
for ci, txt in enumerate(['QTD', 'PRODUTO']):
    set_cell_bg(hrow.cells[ci], '1B3A5C')
    p = hrow.cells[ci].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(txt)
    r.font.name  = 'Calibri'
    r.font.size  = Pt(10)
    r.font.bold  = True
    r.font.color.rgb = WHITE

for ri, (qtd, produto) in enumerate(itens_081):
    row = escopo_tbl.rows[ri + 1]
    bg  = 'E8F1FA' if ri % 2 == 0 else 'FFFFFF'
    set_cell_bg(row.cells[0], bg)
    set_cell_bg(row.cells[1], bg)

    pq = row.cells[0].paragraphs[0]
    pq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pq.paragraph_format.space_before = Pt(4)
    pq.paragraph_format.space_after  = Pt(4)
    rq = pq.add_run(qtd)
    rq.font.name  = 'Calibri'
    rq.font.size  = Pt(11)
    rq.font.bold  = True
    rq.font.color.rgb = NAVY

    pd = row.cells[1].paragraphs[0]
    pd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pd.paragraph_format.space_before = Pt(4)
    pd.paragraph_format.space_after  = Pt(4)
    rd = pd.add_run(produto)
    rd.font.name  = 'Calibri'
    rd.font.size  = Pt(10)
    rd.font.color.rgb = DARK

doc.add_paragraph()

add_heading(doc, '3.2 Dados Técnicos', level=2)
add_kv_table(doc, [
    ('Material de Construção', 'Aço Inox AISI 304'),
    ('Pressão de Operação',    '25 bar'),
    ('Vazão Nominal',          '1,5 L/min por bico'),
    ('Conexão Hidráulica',     'Rosca 1/8" NPT'),
    ('Diâmetro do Orifício',   '0,06 mm'),
    ('Curso de Ajuste',        '150 mm'),
    ('Fluido de Operação',     'Água fresca filtrada'),
])

doc.add_paragraph()
add_heading(doc, '3.3 Características Construtivas', level=2)
add_kv_table(doc, [
    ('Corpo Principal',
     'Construído em perfil tubular de aço inoxidável AISI 304, desenvolvido para suportar '
     'operação contínua em ambiente industrial, proporcionando elevada resistência mecânica '
     'e excelente resistência à corrosão.'),
    ('Sistema de Microposicionamento',
     'Equipado com mecanismo de ajuste linear por fuso roscado protegido, acionado por '
     'manípulo ergonômico com trava mecânica. Permite ajuste fino e repetitivo do '
     'posicionamento do bico de corte, garantindo elevada precisão operacional.'),
    ('Curso Útil de Regulagem', '150 mm'),
    ('Suporte de Posicionamento',
     'Construído em perfil tubular DN 3/4" Schedule 10 em aço inoxidável AISI 304. '
     'Permite regulagem angular do conjunto de corte para otimização da incidência do '
     'jato sobre a folha e facilidade de manutenção.'),
])

doc.add_paragraph()
add_heading(doc, '3.4 Recursos Operacionais', level=2)
for item in [
    'Ajuste fino do posicionamento do bico',
    'Regulagem angular do jato',
    'Facilidade para substituição de bicos',
    'Facilidade de manutenção preventiva e corretiva',
]:
    add_body(doc, item, bullet=True)

add_heading(doc, '3.5 Observações', level=2)
add_body(doc, 'O sistema não contempla ajuste vertical entre tela e bico.')
add_body(doc, 'A regulagem deverá ser executada durante a instalação, considerando distância nominal de 75 mm.')
add_body(doc, 'Itens inclusos:', bold=True)
for item in ['Manômetro de processo', 'Válvula de bloqueio ON/OFF']:
    add_body(doc, item, bullet=True)

add_heading(doc, '3.6 Benefícios Operacionais', level=2)
for item in [
    'Redução de perdas de produção',
    'Melhor estabilidade do corte de refilo',
    'Menor tempo de setup',
    'Maior vida útil dos componentes',
    'Operação segura e confiável',
]:
    add_body(doc, item, bullet=True)

# ── 4. RECOMENDAÇÕES TÉCNICAS ──────────────────────────────────────────────────
add_heading(doc, '4. Recomendações Técnicas')
add_body(doc, 'Recomenda-se o fornecimento de:')
for item in [
    '04 corpos de corte',
    '02 suportes de posicionamento',
    '01 suporte L.A.',
    '01 suporte L.C.',
]:
    add_body(doc, item, bullet=True)
add_body(doc,
    'Para processos com histórico de quebras de papel sem causa definida, recomenda-se a '
    'instalação de desaerador na linha de água, visando eliminar bolhas de ar que possam '
    'comprometer a estabilidade do jato de corte.')

# ── 5. INVESTIMENTO ────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '5. Investimento')

add_investment_block(
    doc,
    'ITEM 01 – SISTEMA DE CORTE DE REFILO',
    [
        ('02 Corpos de Corte', '02', 'R$ 16.950,00', 'R$ 33.900,00'),
        ('02 Suportes',        '02', 'R$  9.900,00', 'R$ 19.800,00'),
    ],
    'TOTAL',
    'R$ 53.700,00'
)

doc.add_paragraph()

# ── 6. PRAZOS E CONDIÇÕES ──────────────────────────────────────────────────────
add_heading(doc, '6. Prazo de Entrega')
add_body(doc, 'Item 01:  60 dias corridos após confirmação do pedido.')

add_heading(doc, '7. Condições de Pagamento')
add_body(doc, 'Item 01:', bold=True)
for item in [
    '15% no pedido',
    '15% na entrega',
    'Saldo em 28/45 DDL',
]:
    add_body(doc, item, bullet=True)

add_heading(doc, '8. Exclusões do Fornecimento')
add_body(doc, 'Não estão inclusos nesta proposta:')
for item in [
    'Transporte e frete',
    'Seguro',
    'Armazenagem',
    'Montagem mecânica',
    'Instalação elétrica',
    'Startup',
    'Supervisão de montagem',
    'Treinamento operacional',
]:
    add_body(doc, item, bullet=True)

add_heading(doc, '9. Garantia')
add_body(doc,
    'Garantia de 180 dias após startup ou 12 meses após faturamento, prevalecendo o '
    'primeiro evento ocorrido.')
add_body(doc,
    'A garantia cobre defeitos comprovados de fabricação, materiais e montagem, desde que '
    'os equipamentos sejam transportados, instalados e operados de acordo com as '
    'recomendações técnicas da TAUBE EQUIPAMENTOS.')

# ── RODAPÉ ─────────────────────────────────────────────────────────────────────
doc.add_paragraph()
sep2 = doc.add_paragraph()
sep2.paragraph_format.space_before = Pt(0)
sep2.paragraph_format.space_after  = Pt(6)
pPr2  = sep2._p.get_or_add_pPr()
pBdr2 = OxmlElement('w:pBdr')
top2  = OxmlElement('w:top')
top2.set(qn('w:val'),   'single')
top2.set(qn('w:sz'),    '6')
top2.set(qn('w:color'), '1B3A5C')
pBdr2.append(top2)
pPr2.append(pBdr2)

footer_tbl = doc.add_table(rows=1, cols=2)
footer_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
no_border(footer_tbl)
fl = footer_tbl.rows[0].cells[0]
fr = footer_tbl.rows[0].cells[1]
fl.width = Cm(11)
fr.width = Cm(6.5)

HYPERLINK_TYPE = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink'

def _gray(para, text, size=9):
    r = para.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(size); r.font.color.rgb = GRAY

def _body_link(para, url, text, size=9):
    rel_id = doc.part.relate_to(url, HYPERLINK_TYPE, is_external=True)
    hl = OxmlElement('w:hyperlink'); hl.set(qn('r:id'), rel_id)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    fn = OxmlElement('w:rFonts')
    fn.set(qn('w:ascii'), 'Calibri'); fn.set(qn('w:hAnsi'), 'Calibri')
    sz  = OxmlElement('w:sz');    sz.set(qn('w:val'), str(size * 2))
    clr = OxmlElement('w:color'); clr.set(qn('w:val'), '555555')
    rPr.extend([fn, sz, clr]); r.append(rPr)
    t = OxmlElement('w:t'); t.text = text; r.append(t)
    hl.append(r); para._p.append(hl)

def _rp(cell, first=False):
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    return p

# Célula esquerda
p_l = fl.paragraphs[0]
p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_l.paragraph_format.space_before = Pt(0)
p_l.paragraph_format.space_after  = Pt(0)
_gray(p_l, 'Aproveitamos para nos colocar à disposição para quaisquer esclarecimentos adicionais.\nAtenciosamente,\n\nTAUBE EQUIPAMENTOS')

# Célula direita — um parágrafo por linha para permitir hyperlinks
_gray(_rp(fr, first=True), 'CNPJ: 11.478.834/0001-23 | IE: 256.047.162')
_gray(_rp(fr), 'Simples Nacional – Não gera crédito de ICMS/IPI')
_body_link(_rp(fr), 'https://www.taubeequipamentos.com.br', 'www.taubeequipamentos.com.br')
p_tel = _rp(fr)
_gray(p_tel, '47 3234-1222 | ')
_body_link(p_tel, 'https://wa.me/5547991661269', '47 9 9166-1269')

build_footer(doc)

out = '/home/henrique/projetos/editor/proposta/PROPOSTA COMERCIAL 081 2026 IRANI SC - TAUBE.docx'
doc.save(out)
print(f'Salvo em: {out}')
