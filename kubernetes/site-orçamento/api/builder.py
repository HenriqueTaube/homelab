"""Geração do .docx de proposta comercial Taube Equipamentos."""
import io
import os
from datetime import date

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

LOGO_PATH = os.path.join(os.path.dirname(__file__), '..', 'assets', 'logo taube minizado.jpg')

NAVY  = RGBColor(0x1B, 0x3A, 0x5C)
STEEL = RGBColor(0x2E, 0x6D, 0xA8)
LIGHT = RGBColor(0xE8, 0xF1, 0xFA)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY  = RGBColor(0x55, 0x55, 0x55)
DARK  = RGBColor(0x22, 0x22, 0x22)

HYPERLINK_TYPE = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink'


# ── helpers XML ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def no_border(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'none')
        el.set(qn('w:sz'),    '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    tblPr.append(tblBorders)

def fix_col_widths(table, widths_dxa):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    for old in tblPr.findall(qn('w:tblW')): tblPr.remove(old)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(sum(widths_dxa)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.insert(0, tblW)
    for old in tblPr.findall(qn('w:tblLayout')): tblPr.remove(old)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    for old in tbl.findall(qn('w:tblGrid')): tbl.remove(old)
    tblGrid = OxmlElement('w:tblGrid')
    for w in widths_dxa:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        tblGrid.append(gc)
    tblPr.addnext(tblGrid)
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for old in tcPr.findall(qn('w:tcW')): tcPr.remove(old)
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(widths_dxa[ci]))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.insert(0, tcW)

# ── helpers de conteúdo ────────────────────────────────────────────────────────

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper() if level == 1 else text)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(13 if level == 1 else 11)
    run.font.bold  = True
    run.font.color.rgb = NAVY
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    btm  = OxmlElement('w:bottom')
    btm.set(qn('w:val'),   'single')
    btm.set(qn('w:sz'),    '6')
    btm.set(qn('w:color'), '1B3A5C')
    pBdr.append(btm)
    pPr.append(pBdr)

def add_body(doc, text, bold=False, bullet=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    if bullet:
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run('• ' + text)
    else:
        run = p.add_run(text)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(10.5)
    run.font.bold  = bold
    run.font.color.rgb = DARK

def add_kv_table(doc, rows, col_widths_dxa=(3575, 6400)):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    no_border(table)
    fix_col_widths(table, list(col_widths_dxa))
    for i, (key, val) in enumerate(rows):
        row = table.rows[i]
        bg  = 'E8F1FA' if i % 2 == 0 else 'FFFFFF'
        for cell, txt, bld in [(row.cells[0], key, True), (row.cells[1], val, False)]:
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent  = Cm(0.2)
            run = p.add_run(txt)
            run.font.name  = 'Calibri'
            run.font.size  = Pt(10)
            run.font.bold  = bld
            run.font.color.rgb = NAVY if bld else DARK

def add_scope_table(doc, itens):
    tbl = doc.add_table(rows=1 + len(itens), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    no_border(tbl)
    fix_col_widths(tbl, [1363, 8562])
    hrow = tbl.rows[0]
    for ci, txt in enumerate(['QTD', 'PRODUTO']):
        set_cell_bg(hrow.cells[ci], '1B3A5C')
        p = hrow.cells[ci].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(txt)
        r.font.name  = 'Calibri'
        r.font.size  = Pt(10)
        r.font.bold  = True
        r.font.color.rgb = WHITE
    for ri, (qtd, produto) in enumerate(itens):
        row = tbl.rows[ri + 1]
        bg  = 'E8F1FA' if ri % 2 == 0 else 'FFFFFF'
        set_cell_bg(row.cells[0], bg)
        set_cell_bg(row.cells[1], bg)
        pq = row.cells[0].paragraphs[0]
        pq.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pq.paragraph_format.space_before = Pt(4)
        pq.paragraph_format.space_after  = Pt(4)
        rq = pq.add_run(qtd)
        rq.font.name  = 'Calibri'
        rq.font.size  = Pt(11)
        rq.font.bold  = True
        rq.font.color.rgb = NAVY
        pd = row.cells[1].paragraphs[0]
        pd.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pd.paragraph_format.space_before = Pt(4)
        pd.paragraph_format.space_after  = Pt(4)
        rd = pd.add_run(produto)
        rd.font.name  = 'Calibri'
        rd.font.size  = Pt(10)
        rd.font.color.rgb = DARK

def add_unified_investment_table(doc, blocks, grand_total):
    # blocks: lista de (nome, qty, valor_unit, total) — uma linha por item
    n_rows = 1 + len(blocks) + 1

    table = doc.add_table(rows=n_rows, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    no_border(table)
    fix_col_widths(table, [5200, 900, 2000, 1875])

    aligns = [
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
    ]
    ri = 0

    for ci, h in enumerate(['Nome do Item', 'Qtd.', 'Valor Unit.', 'Total']):
        set_cell_bg(table.rows[ri].cells[ci], '1B3A5C')
        p = table.rows[ri].cells[ci].paragraphs[0]
        p.alignment = aligns[ci]
        p.paragraph_format.left_indent  = Cm(0.2)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(h)
        run.font.name = 'Calibri'; run.font.size = Pt(10)
        run.font.bold = True; run.font.color.rgb = WHITE
    ri += 1

    for nome, qty, valor_unit, total in blocks:
        for ci, (txt, align) in enumerate(zip([nome, qty, valor_unit, total], aligns)):
            set_cell_bg(table.rows[ri].cells[ci], '2E6DA8')
            p = table.rows[ri].cells[ci].paragraphs[0]
            p.alignment = align
            p.paragraph_format.left_indent  = Cm(0.2)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            run = p.add_run(txt)
            run.font.name = 'Calibri'; run.font.size = Pt(10)
            run.font.bold = True; run.font.color.rgb = WHITE
        ri += 1

    # Total (navy)
    for ci in range(4):
        set_cell_bg(table.rows[ri].cells[ci], '1B3A5C')
        p = table.rows[ri].cells[ci].paragraphs[0]
        p.paragraph_format.left_indent  = Cm(0.2)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        txt = 'TOTAL' if ci == 0 else (grand_total if ci == 3 else '')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 3 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(txt)
        run.font.name = 'Calibri'; run.font.size = Pt(10)
        run.font.bold = True; run.font.color.rgb = WHITE

def build_footer(doc):
    sec = doc.sections[0]
    sec.footer_distance = Cm(0.8)
    footer = sec.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after  = Pt(0)
    pPr = fp._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    for val, pos in [('center', '4819'), ('right', '9638')]:
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), val)
        tab.set(qn('w:pos'), pos)
        tabs.append(tab)
    pPr.append(tabs)

    def make_rPr():
        rPr = OxmlElement('w:rPr')
        fn  = OxmlElement('w:rFonts')
        fn.set(qn('w:ascii'), 'Calibri')
        fn.set(qn('w:hAnsi'), 'Calibri')
        sz  = OxmlElement('w:sz');  sz.set(qn('w:val'), '15')
        clr = OxmlElement('w:color'); clr.set(qn('w:val'), '1B3A5C')
        rPr.extend([fn, sz, clr])
        return rPr

    def add_text(text):
        r = OxmlElement('w:r'); r.append(make_rPr())
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t); fp._p.append(r)

    def add_field(code):
        for ftype in ('begin', 'instr', 'separate', 'end'):
            r = OxmlElement('w:r'); r.append(make_rPr())
            if ftype == 'instr':
                instr = OxmlElement('w:instrText')
                instr.set(qn('xml:space'), 'preserve')
                instr.text = f' {code} '
                r.append(instr)
            else:
                fc = OxmlElement('w:fldChar')
                fc.set(qn('w:fldCharType'), ftype)
                r.append(fc)
            fp._p.append(r)

    def add_hyperlink(url, text):
        rel_id = footer.part.relate_to(url, HYPERLINK_TYPE, is_external=True)
        hl = OxmlElement('w:hyperlink'); hl.set(qn('r:id'), rel_id)
        r = OxmlElement('w:r'); r.append(make_rPr())
        t = OxmlElement('w:t'); t.text = text
        r.append(t); hl.append(r); fp._p.append(hl)

    add_hyperlink('https://wa.me/5547991661269', '47 9 9166-1269')
    add_text('\t'); add_field('PAGE'); add_text(' / '); add_field('NUMPAGES')
    add_text('\t'); add_hyperlink('https://www.taubeequipamentos.com.br', 'taubeequipamentos.com.br')

# ── construção do documento ────────────────────────────────────────────────────

def build_proposta(data: dict) -> bytes:
    """Recebe o payload enriquecido pela IA e retorna os bytes do .docx."""
    cliente      = data['cliente']
    itens        = data['itens']
    investimentos = data['investimentos']
    grand_total  = data.get('grandTotal', 'R$ 0,00')
    apresentacao = data.get('apresentacao', '')
    objetivo     = data.get('objetivo', '')
    prazo        = data.get('prazo', '')
    condicao     = data.get('condicao', '')
    startup      = data.get('startup', 'A consultar')
    num_proposta = data.get('numeroProposta', '000 / 2026')
    hoje         = date.today().strftime('%d/%m/%Y')

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin    = Cm(1.5)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)

    # ── Cabeçalho ──────────────────────────────────────────────────────────────
    header_tbl = doc.add_table(rows=1, cols=2)
    header_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    no_border(header_tbl)
    logo_cell  = header_tbl.rows[0].cells[0]
    title_cell = header_tbl.rows[0].cells[1]
    logo_cell.width  = Cm(5.5)
    title_cell.width = Cm(12.0)
    set_cell_bg(logo_cell,  '1B3A5C')
    set_cell_bg(title_cell, '1B3A5C')
    logo_cell.vertical_alignment  = WD_ALIGN_VERTICAL.CENTER
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    logo_p = logo_cell.paragraphs[0]
    logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_p.paragraph_format.space_before = Pt(6)
    logo_p.paragraph_format.space_after  = Pt(6)
    logo_p.paragraph_format.left_indent  = Cm(0.2)
    logo_p.add_run().add_picture(LOGO_PATH, height=Cm(2.8))
    title_p = title_cell.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_p.paragraph_format.space_before = Pt(10)
    title_p.paragraph_format.space_after  = Pt(2)
    title_p.paragraph_format.right_indent = Cm(0.3)
    r1 = title_p.add_run('TAUBE EQUIPAMENTOS')
    r1.font.name = 'Calibri'; r1.font.size = Pt(16); r1.font.bold = True; r1.font.color.rgb = WHITE
    p2 = title_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.right_indent  = Cm(0.3)
    p2.paragraph_format.space_before  = Pt(2)
    p2.paragraph_format.space_after   = Pt(10)
    r2 = p2.add_run('www.taubeequipamentos.com.br')
    r2.font.name = 'Calibri'; r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0xA8, 0xC8, 0xE8)

    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(0); sep.paragraph_format.space_after = Pt(0)
    pPr = sep._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr'); btm = OxmlElement('w:bottom')
    btm.set(qn('w:val'), 'single'); btm.set(qn('w:sz'), '12'); btm.set(qn('w:color'), '2E6DA8')
    pBdr.append(btm); pPr.append(pBdr)

    num_p = doc.add_paragraph()
    num_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    num_p.paragraph_format.space_before = Pt(8)
    num_p.paragraph_format.space_after  = Pt(8)
    r_num = num_p.add_run(f'PROPOSTA COMERCIAL N° {num_proposta}')
    r_num.font.name = 'Calibri'; r_num.font.size = Pt(13); r_num.font.bold = True; r_num.font.color.rgb = NAVY

    # ── Destinatário ───────────────────────────────────────────────────────────
    dest_tbl = doc.add_table(rows=1, cols=2)
    dest_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    no_border(dest_tbl)
    left_cell  = dest_tbl.rows[0].cells[0]
    right_cell = dest_tbl.rows[0].cells[1]
    left_cell.width  = Cm(9.0)
    right_cell.width = Cm(8.5)
    set_cell_bg(left_cell,  'E8F1FA')
    set_cell_bg(right_cell, 'FFFFFF')

    lp = left_cell.paragraphs[0]
    lp.paragraph_format.left_indent  = Cm(0.3)
    lp.paragraph_format.space_before = Pt(4)
    lp.paragraph_format.space_after  = Pt(2)
    r = lp.add_run('DESTINATÁRIO')
    r.font.name = 'Calibri'; r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = STEEL

    for line, bold in [
        (cliente.get('nome', ''),     True),
        (cliente.get('endereco', ''), False),
        (cliente.get('telefone', ''), False),
        (cliente.get('contato', ''),  False),
        (cliente.get('email', ''),    False),
    ]:
        p = left_cell.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.3)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(line)
        run.font.name = 'Calibri'
        run.font.size = Pt(11 if bold else 10)
        run.font.bold = bold
        run.font.color.rgb = NAVY if bold else DARK

    rp = right_cell.paragraphs[0]
    rp.paragraph_format.left_indent  = Cm(0.3)
    rp.paragraph_format.space_before = Pt(4)
    rp.paragraph_format.space_after  = Pt(2)
    r = rp.add_run('EMISSÃO')
    r.font.name = 'Calibri'; r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = STEEL

    for line in [
        f'Data: {hoje}',
        'Validade: 15 dias a partir da emissão',
        'CNPJ: 11.478.834/0001-23',
        'IE: 256.047.162',
        'Simples Nacional',
    ]:
        p = right_cell.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.3)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(line)
        run.font.name = 'Calibri'; run.font.size = Pt(10); run.font.color.rgb = DARK

    for cell in [left_cell, right_cell]:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(0)

    doc.add_paragraph()

    # ── 1. Apresentação ────────────────────────────────────────────────────────
    add_heading(doc, '1. Apresentação')
    for paragrafo in apresentacao.split('\n'):
        if paragrafo.strip():
            add_body(doc, paragrafo.strip())
    doc.add_paragraph()

    # ── 2. Objetivo ────────────────────────────────────────────────────────────
    add_heading(doc, '2. Objetivo')
    for paragrafo in objetivo.split('\n'):
        if paragrafo.strip():
            add_body(doc, paragrafo.strip())

    # ── Itens ──────────────────────────────────────────────────────────────────
    for idx, item in enumerate(itens):
        num_item  = idx + 1
        sec_base  = num_item + 2  # seções começam em 3
        num_str   = str(num_item).zfill(2)
        titulo    = item.get('titulo', f'Item {num_str}')

        doc.add_page_break()
        add_heading(doc, f'{sec_base}. Item {num_str} – {titulo}')

        add_heading(doc, f'{sec_base}.1 Quantidade', level=2)
        escopo = [(r['qtd'], r['produto']) for r in item.get('escopo', [])]
        add_scope_table(doc, escopo)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        add_heading(doc, f'{sec_base}.2 Dados de Aplicação', level=2)
        dados = [(r['campo'], r['valor']) for r in item.get('dadosAplicacao', [])]
        if dados:
            add_kv_table(doc, dados)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        add_heading(doc, f'{sec_base}.3 Características Construtivas', level=2)
        carac = [(r['componente'], r['descricao']) for r in item.get('caracteristicas', [])]
        if carac:
            add_kv_table(doc, carac)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        add_heading(doc, f'{sec_base}.4 Recursos Operacionais', level=2)
        for rec in item.get('recursosOperacionais', []):
            if rec.strip():
                add_body(doc, rec, bullet=True)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        add_heading(doc, f'{sec_base}.5 Benefícios Operacionais', level=2)
        for ben in item.get('beneficiosOperacionais', []):
            if ben.strip():
                add_body(doc, ben, bullet=True)
        doc.add_paragraph()

    # ── Investimento ───────────────────────────────────────────────────────────
    n = len(itens)
    sec_inv = n + 3
    doc.add_page_break()
    add_heading(doc, f'{sec_inv}. Investimento')

    blocks = []
    for idx, (item, inv) in enumerate(zip(itens, investimentos)):
        num_str  = str(idx + 1).zfill(2)
        titulo   = item.get('titulo', f'Item {num_str}')
        qtd      = inv.get('qtd', '01')
        val_unit = inv.get('valorUnit', 'R$ 0,00')
        try:
            import re
            n_val = float(re.sub(r'[R$\s.]', '', val_unit).replace(',', '.'))
            n_qtd = float(qtd)
            total = f"R$ {n_val * n_qtd:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
        except Exception:
            total = val_unit
        blocks.append((titulo, qtd, val_unit, total))

    add_unified_investment_table(doc, blocks, grand_total)
    doc.add_paragraph()

    # ── Prazo, Condições, Exclusões, Garantia, Validade ───────────────────────
    s = sec_inv + 1
    add_heading(doc, f'{s}. Prazo de Entrega')
    for linha in prazo.split('\n'):
        if linha.strip():
            add_body(doc, linha.strip())

    s += 1
    add_heading(doc, f'{s}. Condições de Pagamento')
    for linha in condicao.split('\n'):
        if linha.strip():
            add_body(doc, linha.strip(), bullet=True)
    add_body(doc, 'Obs.: quando houver sinal de entrada, o prazo de entrega passa a ser contado a partir da data de confirmação do recebimento do sinal.')

    s += 1
    add_heading(doc, f'{s}. Dados Complementares')
    add_kv_table(doc, [
        ('Frete',                    'FOB'),
        ('Impostos',                 'Simples Nacional'),
        ('Acompanhamento Startup',   startup or 'A consultar'),
        ('Montagem',                 'Excluída do escopo do fornecimento'),
    ])

    s += 1
    add_heading(doc, f'{s}. Garantia')
    add_body(doc, 'A Taube Equipamentos garante os equipamentos fornecidos pelo prazo de 12 (doze) meses, contados a partir da data de entrega ou entrada em operação, contra defeitos de fabricação e/ou falhas comprovadas de materiais empregados em sua construção.')
    add_body(doc, 'A garantia está limitada ao reparo ou substituição dos componentes reconhecidamente defeituosos, não abrangendo danos decorrentes de utilização inadequada, operação incorreta, negligência, modificações não autorizadas, desgaste natural, falta de manutenção ou quaisquer condições operacionais em desacordo com as recomendações técnicas da Taube Equipamentos.')
    add_body(doc, 'Os reparos em garantia serão realizados com o material posto na fábrica da Taube, sendo de responsabilidade do cliente os custos de embalagem, transporte, frete de envio e retorno dos equipamentos ou componentes a serem avaliados e reparados.')
    add_body(doc, 'Para equipamentos instalados fora do território brasileiro, a garantia permanece válida quanto ao fornecimento de peças e mão de obra para correção dos defeitos cobertos. Entretanto, os custos relacionados à logística internacional, transporte de materiais, deslocamento de técnicos, passagens aéreas, hospedagem, alimentação, translados, taxas locais e demais despesas de viagem não estão contemplados na garantia e serão integralmente de responsabilidade do cliente.')
    add_body(doc, 'Nos atendimentos realizados em campo, a cobertura da garantia restringe-se exclusivamente à mão de obra técnica necessária para execução dos reparos reconhecidos pela Taube Equipamentos como cobertos pelas condições de garantia.')

    s += 1
    add_heading(doc, f'{s}. Validade da Proposta')
    add_body(doc, 'Esta proposta possui validade de 15 dias a partir da data de emissão.')

    # ── Rodapé do documento ────────────────────────────────────────────────────
    doc.add_paragraph()
    sep2 = doc.add_paragraph()
    sep2.paragraph_format.space_before = Pt(0); sep2.paragraph_format.space_after = Pt(6)
    pPr2 = sep2._p.get_or_add_pPr(); pBdr2 = OxmlElement('w:pBdr'); top2 = OxmlElement('w:top')
    top2.set(qn('w:val'), 'single'); top2.set(qn('w:sz'), '6'); top2.set(qn('w:color'), '1B3A5C')
    pBdr2.append(top2); pPr2.append(pBdr2)

    footer_tbl = doc.add_table(rows=1, cols=2)
    footer_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    no_border(footer_tbl)
    fl = footer_tbl.rows[0].cells[0]
    fr = footer_tbl.rows[0].cells[1]
    fl.width = Cm(11); fr.width = Cm(6.5)

    def _gray(para, text):
        r = para.add_run(text)
        r.font.name = 'Calibri'; r.font.size = Pt(9); r.font.color.rgb = GRAY

    def _link(para, url, text):
        rel_id = doc.part.relate_to(url, HYPERLINK_TYPE, is_external=True)
        hl = OxmlElement('w:hyperlink'); hl.set(qn('r:id'), rel_id)
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        fn = OxmlElement('w:rFonts'); fn.set(qn('w:ascii'), 'Calibri'); fn.set(qn('w:hAnsi'), 'Calibri')
        sz = OxmlElement('w:sz');  sz.set(qn('w:val'), '18')
        clr = OxmlElement('w:color'); clr.set(qn('w:val'), '555555')
        rPr.extend([fn, sz, clr]); r.append(rPr)
        t = OxmlElement('w:t'); t.text = text; r.append(t)
        hl.append(r); para._p.append(hl)

    def rp(cell, first=False):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        return p

    p_l = fl.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_l.paragraph_format.space_before = Pt(0); p_l.paragraph_format.space_after = Pt(0)
    _gray(p_l, 'Aproveitamos para nos colocar à disposição para quaisquer esclarecimentos adicionais.\nAtenciosamente,\n\nTAUBE EQUIPAMENTOS')

    _gray(rp(fr, first=True), 'CNPJ: 11.478.834/0001-23 | IE: 256.047.162')
    _gray(rp(fr), 'Simples Nacional – Não gera crédito de ICMS/IPI')
    _link(rp(fr), 'https://www.taubeequipamentos.com.br', 'www.taubeequipamentos.com.br')
    p_tel = rp(fr)
    _gray(p_tel, '47 3234-1222 | ')
    _link(p_tel, 'https://wa.me/5547991661269', '47 9 9166-1269')

    build_footer(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
